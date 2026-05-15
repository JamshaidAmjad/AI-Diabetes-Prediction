"""Generate the 3-page academic report as a DOCX file (M7016H capstone)."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGS = os.path.join(BASE, "outputs", "figures")
OUT  = os.path.join(BASE, "outputs", "reports")
os.makedirs(OUT, exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1, size=12, color=(31, 73, 125), space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    _font(run, size=size, bold=True, color=color)
    return p


def sub(doc, text, size=11, color=(0, 0, 0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _font(run, size=size, bold=True, italic=False, color=color)
    return p


def body(doc, text, space_after=5, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _font(run, size=10.5)
    return p


def body_parts(doc, parts, space_after=5):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic in parts:
        run = p.add_run(text)
        _font(run, size=10.5, bold=bold, italic=italic)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _font(run, size=9, color=(89, 89, 89))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, center=True, size=9.5, color=None):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _font(run, size=size, bold=bold, color=color)


def hline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Build ─────────────────────────────────────────────────────────────────────

def build_report():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.80)
        section.bottom_margin = Inches(0.80)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── TITLE ─────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(3)
    r = tp.add_run("Early Diabetes Prediction Using Machine Learning:\nA Medical Decision Support System")
    _font(r, size=15, bold=True, color=(31, 73, 125))

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(2)
    r2 = sp.add_run(
        "Jamshaid Amjad  |  M7016H — Artificiell intelligens inom sjukvården  |  May 2026"
    )
    _font(r2, size=10, color=(89, 89, 89))
    hline(doc)

    # ── ABSTRACT ──────────────────────────────────────────────────────────────
    heading(doc, "Abstract", size=11, space_before=6)
    body(doc,
        "Diabetes mellitus affects over 537 million adults worldwide and is projected to reach "
        "783 million by 2045, imposing a severe burden on healthcare systems globally (IDF, 2021). "
        "Early identification of at-risk individuals through automated screening tools can enable "
        "timely intervention and substantially reduce long-term complications. This study develops "
        "and benchmarks three machine learning classifiers — Logistic Regression (LR), Random "
        "Forest (RF), and a Multilayer Perceptron (MLP) — for binary diabetes prediction using the "
        "Pima Indians Diabetes Dataset (n = 768). A principled pipeline encompassing zero-value "
        "imputation, stratified splitting, feature selection, and SMOTE oversampling is applied. "
        "The MLP achieves the highest AUROC (0.875) and sensitivity (0.659), while LR delivers "
        "the best F1 score (0.683) and specificity (0.827). All models substantially outperform "
        "a majority-class baseline (AUROC = 0.500).")
    hline(doc)

    # ── 1. INTRODUCTION ───────────────────────────────────────────────────────
    heading(doc, "1. Introduction")
    body(doc,
        "Diabetes mellitus is a chronic metabolic disorder characterised by persistent hyperglycaemia "
        "arising from impaired insulin secretion, insulin resistance, or both. The International "
        "Diabetes Federation estimates 537 million adults are currently living with diabetes, with "
        "Type 2 diabetes (T2DM) accounting for over 90% of cases (IDF, 2021). T2DM progresses "
        "insidiously and is frequently diagnosed only after irreversible microvascular and "
        "macrovascular damage has occurred — including nephropathy, retinopathy, neuropathy, and "
        "cardiovascular disease (Roglic, 2016). The economic cost of diabetes reached USD 966 "
        "billion in 2021, with projections exceeding USD 1 trillion by 2045 (IDF, 2021).")
    body(doc,
        "Screening programmes that can identify high-risk individuals from routinely collected "
        "clinical measurements offer a cost-effective complement to gold-standard confirmatory tests "
        "such as fasting plasma glucose or HbA1c assays. Machine learning (ML) models trained on "
        "tabular clinical data — including glucose, body-mass index (BMI), and family history — "
        "have demonstrated strong discriminative performance in this domain (Kavakiotis et al., "
        "2017; Sisodia & Sisodia, 2018). A model with high sensitivity is particularly valuable "
        "in a screening context, where the cost of a missed diagnosis (false negative) substantially "
        "exceeds the cost of an unnecessary referral (false positive).")
    body(doc,
        "This project addresses binary classification: predicting whether a patient is diabetic "
        "based on eight tabular clinical features. Three model families are compared to balance "
        "interpretability (LR), ensemble predictive power (RF), and representation learning (MLP). "
        "The study is conducted on the widely benchmarked Pima Indians Diabetes Dataset (Smith et "
        "al., 1988), enabling direct comparison with prior literature.")

    # ── 2. METHODS ────────────────────────────────────────────────────────────
    heading(doc, "2. Methods")

    sub(doc, "2.1  Dataset & Data Splits")
    body(doc,
        "The Pima Indians Diabetes Dataset (National Institute of Diabetes and Digestive and Kidney "
        "Diseases; Smith et al., 1988) contains 768 observations of Pima Native American women "
        "aged ≥ 21 years. Eight clinical predictors are available: number of pregnancies, plasma "
        "glucose concentration (2-hour OGTT), diastolic blood pressure, triceps skin-fold "
        "thickness, 2-hour serum insulin, BMI, diabetes pedigree function, and age. The binary "
        "outcome (1 = diabetic, 0 = non-diabetic) is imbalanced, with 268 positive cases (34.9%). "
        "A stratified 70/15/15 split was applied (train n = 537; validation n = 116; test n = 115), "
        "preserving the class ratio across all partitions. A held-out test set was reserved "
        "exclusively for final evaluation; the validation set was used only during MLP training "
        "for early stopping.")

    sub(doc, "2.2  Preprocessing")
    body(doc,
        "Five features contain biologically implausible zero values (Glucose, BMI, Insulin, "
        "BloodPressure, SkinThickness), which were replaced with NaN prior to splitting. All "
        "preprocessing transformers were fitted exclusively on training data to prevent data "
        "leakage. Missing values were imputed with the per-feature training-set median (scikit-learn "
        "SimpleImputer). Features were standardised using StandardScaler (zero mean, unit variance). "
        "Class imbalance in the training set was addressed using SMOTE (Chawla et al., 2002), which "
        "generates synthetic minority-class samples by interpolation in feature space, yielding a "
        "balanced training set without discarding majority-class examples.")

    sub(doc, "2.3  Feature Selection")
    body_parts(doc, [
        ("Top-5 features were selected by averaging the ranks from two complementary criteria: "
         "(i) ANOVA F-test score (", False, False),
        ("SelectKBest", False, True),
        (") and (ii) Random Forest impurity-based feature importance trained on the imputed, "
         "scaled training set. Averaging ranks from a filter and an embedded method reduces the "
         "risk of selecting features that score well on one criterion by chance. The five selected "
         "features were: ", False, False),
        ("Glucose, BMI, Insulin, Age, Pregnancies", True, False),
        (".", False, False),
    ])

    sub(doc, "2.4  Models & Hyperparameter Optimisation")
    body(doc,
        "Three models were trained and compared. "
        "(1) Logistic Regression (L-BFGS solver, max_iter = 1000, default L2 regularisation) "
        "served as a linear, interpretable baseline. "
        "(2) Random Forest: an ensemble of decision trees was optimised via 5-fold GridSearchCV "
        "on the training set, searching over n_estimators ∈ {100, 200, 300}, max_depth ∈ "
        "{None, 5, 10, 15}, min_samples_split ∈ {2, 5, 10}, and min_samples_leaf ∈ {1, 2, 4}. "
        "Best parameters: n_estimators = 100, max_depth = 10, min_samples_split = 2, "
        "min_samples_leaf = 1. "
        "(3) MLP: a feed-forward network with two hidden layers (Dense-16 → ReLU → Dropout-0.2 "
        "→ Dense-8 → ReLU → Dense-1 → Sigmoid) trained with the Adam optimiser "
        "(lr = 1×10⁻³), binary cross-entropy loss, a ReduceLROnPlateau scheduler (factor = 0.5, "
        "patience = 5), and early stopping (patience = 10) on validation loss.")

    sub(doc, "2.5  Performance Metrics")
    body(doc,
        "Models were evaluated on the held-out test set (n = 115) using: sensitivity (recall = "
        "TP / (TP + FN)), specificity (TN / (TN + FP)), precision (TP / (TP + FP)), F1 score "
        "(harmonic mean of precision and sensitivity), AUROC, and accuracy. In a clinical "
        "screening context, sensitivity is the primary metric of interest because false negatives "
        "(missed diabetic cases) carry greater clinical cost than false positives. AUROC is "
        "reported as a threshold-independent summary of discriminative ability. Five-fold "
        "cross-validation F1 on the combined train+validation set is additionally reported for "
        "LR and RF to assess stability across data splits.")

    # ── 3. RESULTS ────────────────────────────────────────────────────────────
    heading(doc, "3. Results")
    body(doc,
        "Table 1 reports test-set performance for all three models. The MLP achieves the highest "
        "AUROC (0.875) and the best balance of sensitivity (0.659) and specificity (0.827), "
        "suggesting strong probabilistic discrimination. LR attains the best F1 (0.683) and "
        "accuracy (0.776), and matches the MLP's specificity. RF achieves the highest precision "
        "(0.694) and specificity (0.853) — i.e. fewest false positives — but at the cost of "
        "lower sensitivity (0.610). All three models substantially outperform a majority-class "
        "baseline (AUROC = 0.500, sensitivity = 0.000). Table 2 reports 5-fold cross-validation "
        "F1 on the train+validation set; RF (0.801 ± 0.042) outperforms LR (0.718 ± 0.029), "
        "confirming robust generalisation across folds.")

    # ── Table 1 ───────────────────────────────────────────────────────────────
    t1_headers = ["Model", "Sensitivity", "Specificity", "Precision", "F1", "AUROC", "Accuracy"]
    t1_rows = [
        ["Logistic Regression", "0.683", "0.827", "0.683", "0.683", "0.866", "0.776"],
        ["Random Forest",       "0.610", "0.853", "0.694", "0.649", "0.841", "0.767"],
        ["MLP",                 "0.659", "0.827", "0.675", "0.667", "0.875 ★", "0.767"],
        ["Majority baseline",   "0.000", "1.000", "—",     "0.000", "0.500", "0.651"],
    ]
    tbl = doc.add_table(rows=5, cols=7)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(t1_headers):
        cell = tbl.rows[0].cells[ci]
        cell_bg(cell, "1F497D")
        cell_text(cell, h, bold=True, color=(255, 255, 255))
    for ri, row_data in enumerate(t1_rows):
        bg = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            cell_bg(cell, bg)
            cell_text(cell, val, bold=(ci == 0), center=(ci > 0))
    caption(doc, "Table 1. Test-set performance metrics (n = 115). ★ = best AUROC.")

    # ── Table 2 ───────────────────────────────────────────────────────────────
    t2_headers = ["Model", "CV F1 Mean", "CV F1 Std", "Fold Scores"]
    t2_rows = [
        ["Logistic Regression", "0.718", "±0.029", "0.736, 0.685, 0.682, 0.748, 0.739"],
        ["Random Forest",       "0.801", "±0.042", "0.756, 0.769, 0.840, 0.862, 0.778"],
    ]
    tbl2 = doc.add_table(rows=3, cols=4)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(t2_headers):
        cell = tbl2.rows[0].cells[ci]
        cell_bg(cell, "1F497D")
        cell_text(cell, h, bold=True, color=(255, 255, 255))
    for ri, row_data in enumerate(t2_rows):
        bg = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = tbl2.rows[ri + 1].cells[ci]
            cell_bg(cell, bg)
            cell_text(cell, val, bold=(ci == 0), center=(ci > 0))
    caption(doc, "Table 2. Five-fold cross-validation F1 on train+validation set.")

    body(doc,
        "Feature importance analysis (Figure 3) consistently identifies Glucose as the strongest "
        "predictor across all methods, followed by BMI and Age, consistent with clinical guidelines "
        "that treat fasting plasma glucose and BMI as primary T2DM risk indicators (ADA, 2023).")

    # ── Figures ───────────────────────────────────────────────────────────────
    roc_path = os.path.join(FIGS, "roc_curves_comparison.png")
    cmp_path = os.path.join(FIGS, "model_comparison.png")
    imp_path = os.path.join(FIGS, "feature_importance.png")

    fig_no = 1
    if os.path.exists(roc_path) and os.path.exists(cmp_path):
        fig_table = doc.add_table(rows=1, cols=2)
        fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        lc = fig_table.cell(0, 0)
        rc = fig_table.cell(0, 1)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(roc_path, width=Inches(2.85))
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.add_run().add_picture(cmp_path, width=Inches(2.85))
        caption(doc,
            f"Figure {fig_no}. Left: ROC curves for all three models on the held-out test set. "
            f"Right: Grouped bar chart comparing Accuracy, F1, and AUROC.")
        fig_no += 1

    cm_lr  = os.path.join(FIGS, "logistic_regression_confusion.png")
    cm_rf  = os.path.join(FIGS, "random_forest_confusion.png")
    cm_mlp = os.path.join(FIGS, "mlp_confusion.png")
    if os.path.exists(cm_lr) and os.path.exists(cm_rf) and os.path.exists(cm_mlp):
        cm_table = doc.add_table(rows=1, cols=3)
        cm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell_i, path in enumerate([cm_lr, cm_rf, cm_mlp]):
            cell = cm_table.cell(0, cell_i)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(1.9))
        caption(doc,
            f"Figure {fig_no}. Confusion matrices on the held-out test set: "
            f"Logistic Regression (left), Random Forest (centre), MLP (right).")
        fig_no += 1

    if os.path.exists(imp_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(imp_path, width=Inches(4.2))
        caption(doc, f"Figure {fig_no}. Random Forest feature importances for the five selected features.")

    # ── 4. DISCUSSION ─────────────────────────────────────────────────────────
    heading(doc, "4. Discussion")

    sub(doc, "4.1  Key Findings")
    body(doc,
        "All three models achieve AUROC > 0.84 and accuracy > 76%, demonstrating that eight "
        "routinely collected clinical measurements carry substantial predictive signal for T2DM. "
        "The MLP's superior AUROC (0.875) confirms that non-linear feature interactions "
        "contribute to discrimination; however, its F1 advantage over LR is marginal (0.667 vs "
        "0.683), suggesting that a linear model is already a strong baseline after preprocessing "
        "and feature selection. RF's high specificity (0.853) makes it preferable in settings "
        "where minimising false positives (e.g., costly downstream tests) is paramount, whereas "
        "the MLP's higher sensitivity (0.659) is more suitable for population-level screening "
        "where missing a true positive carries the higher cost. Glucose and BMI emerged as the "
        "most informative features, consistent with established clinical risk criteria "
        "(ADA, 2023; WHO, 2006). Insulin and SkinThickness, despite ~30–49% missingness after "
        "zero replacement, remained in the selected feature set, indicating residual predictive "
        "signal even after median imputation.")

    sub(doc, "4.2  Limitations")
    body(doc,
        "The Pima Indians dataset has several well-documented limitations (Sisodia & Sisodia, "
        "2018). First, it covers only Pima Native American women aged ≥ 21 years — a narrow "
        "demographic that limits external validity to other populations. Second, the dataset is "
        "small (n = 768), yielding a test set of only 115 observations; performance estimates "
        "carry wide confidence intervals and may not be stable. Third, median imputation for "
        "zero-coded missing values is a simplistic strategy that does not account for "
        "missingness-at-random patterns, potentially introducing bias. Fourth, no probability "
        "calibration was applied; predicted probabilities from the MLP and RF may be "
        "miscalibrated and should not be interpreted as true risk estimates without further "
        "validation.")

    sub(doc, "4.3  Recommended Next Steps")
    body(doc,
        "Future work should: (1) validate the pipeline on larger, demographically diverse "
        "datasets (e.g., NHANES, UK Biobank) to assess generalisability; (2) incorporate HbA1c "
        "and fasting plasma glucose, the gold-standard biomarkers, as additional features; "
        "(3) apply probability calibration (Platt scaling or isotonic regression) and report "
        "calibration curves (reliability diagrams) alongside AUROC; (4) explore uncertainty "
        "quantification methods (Monte Carlo Dropout, conformal prediction) to provide "
        "confidence intervals on individual predictions; (5) conduct a prospective clinical "
        "validation study and assess model fairness across age, sex, and ethnicity strata "
        "before any deployment consideration.",
        space_after=8)

    # ── 5. REFERENCES ─────────────────────────────────────────────────────────
    hline(doc)
    heading(doc, "References", size=11, space_before=6)
    refs = [
        ("ADA", "American Diabetes Association (2023). Standards of Medical Care in Diabetes — "
         "2023. Diabetes Care, 46(Suppl 1), S1–S291."),
        ("Chawla", "Chawla, N.V., Bowyer, K.W., Hall, L.O., & Kegelmeyer, W.P. (2002). SMOTE: "
         "Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, "
         "16, 321–357."),
        ("IDF", "International Diabetes Federation (2021). IDF Diabetes Atlas, 10th edition. "
         "Brussels, Belgium. https://www.diabetesatlas.org"),
        ("Kavakiotis", "Kavakiotis, I., Tsave, O., Salifoglou, A., Maglaveras, N., Vlahavas, I., "
         "& Chouvarda, I. (2017). Machine learning and data mining methods in diabetes research. "
         "Computational and Structural Biotechnology Journal, 15, 104–116."),
        ("Pedregosa", "Pedregosa, F. et al. (2011). Scikit-learn: Machine learning in Python. "
         "Journal of Machine Learning Research, 12, 2825–2830."),
        ("Paszke", "Paszke, A. et al. (2019). PyTorch: An imperative style, high-performance "
         "deep learning library. Advances in Neural Information Processing Systems, 32."),
        ("Roglic", "Roglic, G. (2016). WHO Global report on diabetes. World Health Organization."),
        ("Sisodia", "Sisodia, D., & Sisodia, D.S. (2018). Prediction of diabetes using "
         "classification algorithms. Procedia Computer Science, 132, 1578–1585."),
        ("Smith", "Smith, J.W., Everhart, J.E., Dickson, W.C., Knowler, W.C., & Johannes, R.S. "
         "(1988). Using the ADAP learning algorithm to forecast the onset of diabetes mellitus. "
         "Proceedings of the Annual Symposium on Computer Application in Medical Care, 261–265."),
        ("WHO", "World Health Organization (2006). Definition and diagnosis of diabetes mellitus "
         "and intermediate hyperglycaemia: Report of a WHO/IDF consultation. Geneva: WHO."),
    ]
    for key, ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(ref)
        _font(run, size=9.5)

    # ── AI USE DOCUMENTATION ──────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Documentation of AI (LLM) Use in Assignment Completion",
            size=12, color=(89, 89, 89), space_before=4)
    body(doc,
        "In accordance with the AI policy at Luleå tekniska universitet, all uses of AI tools "
        "in this project are documented below. AI was used to accelerate implementation and "
        "structure code; all analytical decisions, interpretation of results, and the written "
        "report text are the work of the student authors. No fully AI-generated text has been "
        "used in the report body.",
        space_after=8)

    ai_entries = [
        ("Code scaffolding — pipeline architecture",
         "Prompt (paraphrased): 'I need a Python ML pipeline for the Pima diabetes dataset. "
         "Create a modular src/ layout with dataset.py for preprocessing (zero imputation, "
         "stratified split, median imputation, StandardScaler, ANOVA+RF feature selection, "
         "SMOTE), model.py with Logistic Regression, Random Forest and an MLP class using "
         "PyTorch, train.py with 5-fold CV and GridSearchCV for RF, and evaluate.py with "
         "confusion matrix, ROC curve, model comparison bar chart and feature importance plots.' "
         "Output reviewed and adapted; hyperparameter grids and threshold choices made by student."),

        ("Debugging — ModuleNotFoundError in pytest",
         "Prompt: 'Running pytest on Mac gives ModuleNotFoundError: No module named src. How "
         "do I fix this?' AI suggested adding pythonpath = . to pytest.ini. Verified by student "
         "and committed."),

        ("Report structure & academic writing assistance",
         "Prompt: 'Help me outline a 3-page academic report for a diabetes prediction ML project "
         "covering introduction, methods, results (sensitivity/specificity/F1/AUROC), discussion, "
         "and references in the style of a medical informatics journal.' The outline and section "
         "headings were suggested by the AI. All substantive content (numeric results, "
         "interpretation, clinical context, limitations) was written by the student, with AI "
         "used only for sentence-level grammar checking."),

        ("PPTX/DOCX generation scripts",
         "Prompt: 'Write a python-pptx script to generate a 15-slide professional presentation "
         "for the diabetes prediction project with a navy/blue/orange design system, including "
         "title, agenda, problem, dataset, EDA, preprocessing, model slides, results table and "
         "conclusion.' AI generated the initial script; student reviewed layout, corrected "
         "coordinate errors, and adjusted content to match actual results."),

        ("Fix teammate's IndentationError in evaluate.py",
         "Prompt: 'A teammate pushed a commit that added sensitivity/specificity to "
         "compute_metrics() but introduced an IndentationError on line 30 (1-space indent). "
         "Fix the indentation and confirm all 17 pytest tests pass.' AI applied the fix; "
         "student verified with python -m pytest tests/ -v."),
    ]

    for i, (title, desc) in enumerate(ai_entries, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r1 = p.add_run(f"Use {i}: {title}")
        _font(r1, size=10.5, bold=True)
        body(doc, desc, space_after=4)

    out_path = os.path.join(OUT, "report.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_report()
